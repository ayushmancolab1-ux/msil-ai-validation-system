"""
extractor.py
Extracts structured engineering fields from uploaded documents:
  - Image files (PNG/JPG/WEBP) → LLM vision agent (Azure GPT-4.1 or OpenAI)
  - PDF files → pdfplumber text extraction + regex
  - DOCX files → python-docx text extraction + regex
Falls back to seeded JSON data when extraction yields insufficient fields.
"""
import re
import json
import logging
from pathlib import Path

import pdfplumber
from docx import Document

from engine.llm_agent import (
    extract_drawing_from_image,
    is_image_file,
    IMAGE_EXTENSIONS,
)

logger = logging.getLogger(__name__)

SEEDED_DATA_DIR = Path(__file__).parent.parent / "data"


def extract_fields(
    filepath: str,
    vehicle_model: str,
    component: str,
    doc_type: str = "drawing",
    llm_provider: str = "openai",
    plant: str = "Manesar",
) -> dict:
    """
    Extract structured fields from a drawing or WIS document.

    For drawing images → uses LLM vision.
    For PDF/DOCX → uses regex extraction.
    Always merges with seeded data to fill gaps.
    """
    ext = Path(filepath).suffix.lower()
    extracted = {}

    try:
        if ext in IMAGE_EXTENSIONS and doc_type == "drawing":
            # ── LLM vision path ──────────────────────────────────────────────
            extracted = extract_drawing_from_image(filepath, provider=llm_provider)
        elif ext == ".pdf":
            text = _extract_pdf_text(filepath)
            extracted = _regex_extract(text)
        elif ext in {".docx", ".doc"}:
            text = _extract_docx_text(filepath)
            extracted = _regex_extract(text)
        else:
            logger.warning("Unsupported file extension %s — using seeded data only", ext)
    except Exception as e:
        logger.error("Extraction failed for %s: %s", filepath, e)
        extracted = {}

    # Load seeded data as baseline (always present)
    seeded = _load_seeded_data(vehicle_model, plant, component, doc_type)

    # Merge: seeded fills missing fields; extracted overrides where present and non-empty
    result = {**seeded}
    for k, v in extracted.items():
        if v is not None and v != "" and v != []:
            result[k] = v

    return result


# ─────────────────────────────────────────────────────────────────────────────
# Text extraction helpers
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf_text(filepath: str) -> str:
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    except Exception as e:
        logger.warning("PDF extraction error: %s", e)
    return text


def _extract_docx_text(filepath: str) -> str:
    text = ""
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
    except Exception as e:
        logger.warning("DOCX extraction error: %s", e)
    return text


# ─────────────────────────────────────────────────────────────────────────────
# Regex extraction (PDF / DOCX path)
# ─────────────────────────────────────────────────────────────────────────────

# Fields that only exist on a Drawing, never in a WIS text document.
# We strip these from the WIS seeded fallback so they don't cause false deviations.
_WIS_EXCLUDE_FIELDS = {
    "bore_diameter_mm",
    "overall_width_mm",
    "overall_height_mm",
    "shaft_length_mm",
}


def _regex_extract(text: str) -> dict:
    result = {}

    # Torque: collect ALL "X Nm" values and use the maximum (final torque in multi-stage sequences)
    torque_matches = re.findall(r'(\d+\.?\d*)\s*Nm', text, re.IGNORECASE)
    if torque_matches:
        result["primary_torque_nm"] = max(float(v) for v in torque_matches)

    # Bore diameter: ⌀64.00 or ø64.00
    bore_match = re.search(r'[⌀Øø](\d+\.?\d*)', text)
    if bore_match:
        result["bore_diameter_mm"] = float(bore_match.group(1))

    # First ± value → bore tolerance
    tol_matches = re.findall(r'±(\d+\.?\d*)', text)
    if tol_matches:
        result["bore_tolerance_mm"] = float(tol_matches[0])
    if len(tol_matches) > 1:
        result["gap_tolerance_mm"] = float(tol_matches[1])

    # Ra value
    ra_match = re.search(r'Ra\s*(\d+\.?\d*)', text, re.IGNORECASE)
    if ra_match:
        result["Ra_value"] = float(ra_match.group(1))

    # Revision mark
    rev_match = re.search(r'\bREV\b\.?\s*([A-Z])', text, re.IGNORECASE)
    if rev_match:
        result["drawing_revision"] = rev_match.group(1).upper()

    # Drawing number MSIL-XX-XX-XXXX-X
    pn_match = re.search(r'(MSIL-[A-Z]{2}-[A-Z]{2}-\d{4}-[A-Z])', text)
    if pn_match:
        result["drawing_number"] = pn_match.group(1)

    # Part number SW-XXX-XXX-000-X
    part_match = re.search(r'([A-Z]{2}-[A-Z]{2,4}-[A-Z]{2,5}-\d{3}-[A-Z])', text)
    if part_match:
        result["part_number"] = part_match.group(1)

    # Fastener size
    fastener_match = re.search(r'\b(M\d+)\b', text)
    if fastener_match:
        result["fastener_size"] = fastener_match.group(1)

    # Fastener grade e.g. 8.8 or 10.9
    grade_match = re.search(r'\b(8\.8|10\.9|12\.9|4\.8|6\.8)\b', text)
    if grade_match:
        result["fastener_grade"] = grade_match.group(1)

    # Flatness tolerance
    flat_match = re.search(r'(?:flatness|□)\s*(\d+\.?\d*)', text, re.IGNORECASE)
    if flat_match:
        result["flatness_tolerance_mm"] = float(flat_match.group(1))

    # Parallelism
    para_match = re.search(r'(?:parallelism|//)\s*(\d+\.?\d*)', text, re.IGNORECASE)
    if para_match:
        result["parallelism_tolerance_mm"] = float(para_match.group(1))

    # True position
    pos_match = re.search(r'(?:true.?position|⊕|position)\s*[⌀Øø]?\s*(\d+\.?\d*)', text, re.IGNORECASE)
    if pos_match:
        result["true_position_dia_mm"] = float(pos_match.group(1))

    # Thickness t = 4.0
    thick_match = re.search(r't\s*=\s*(\d+\.?\d*)', text, re.IGNORECASE)
    if thick_match:
        result["thickness_mm"] = float(thick_match.group(1))

    # Material grade  IS 2062 Gr. B
    mat_match = re.search(r'(IS\s*\d+\s*Gr\.?\s*[A-Z])', text, re.IGNORECASE)
    if mat_match:
        result["material_grade"] = mat_match.group(1).strip()

    # Heat treatment
    ht_match = re.search(r'HEAT\s*TREATMENT[:\s]*([A-Z]+)', text, re.IGNORECASE)
    if ht_match:
        result["heat_treatment"] = ht_match.group(1).strip()

    # Datum references
    datum_matches = re.findall(r'\bDATUM\s+([A-Z])\b', text, re.IGNORECASE)
    if datum_matches:
        result["datum_references"] = list(dict.fromkeys(datum_matches))

    return result


# ─────────────────────────────────────────────────────────────────────────────
# Seeded data loader
# ─────────────────────────────────────────────────────────────────────────────

_MODEL_MAP = {
    "Maruti Swift": "swift",
    "Maruti Brezza": "brezza",
    "Maruti Dzire": "dzire",
}

_PLANT_MAP = {
    "Manesar": "manesar",
    "Gurugram": "gurugram",
    "Gujarat": "gujarat",
}

_COMP_MAP = {
    "Engine Mount Bracket (Front LH)": "engine_mount",
    "Brake Caliper Bolt Assembly": "brake_caliper",
    "Suspension Strut Upper Mount": "suspension_strut",
    "Steering Rack Mounting Bracket": "steering_rack",
    "Fuel Tank Strap Assembly": "fuel_tank",
    "Exhaust Manifold Stud": "exhaust_manifold",
    "Wheel Hub Bearing": "wheel_hub",
    "Gearbox Crossmember": "gearbox",
    # Legacy names — map to nearest
    "Engine Assembly": "engine_mount",
    "Transmission Assembly": "gearbox",
    "Brake System": "brake_caliper",
    "Suspension System": "suspension_strut",
    "Steering Assembly": "steering_rack",
    "Fuel System": "fuel_tank",
    "Exhaust System": "exhaust_manifold",
    "Electrical Harness": "engine_mount",
}


def _load_seeded_data(vehicle_model: str, plant: str, component: str, doc_type: str) -> dict:
    model_key = _MODEL_MAP.get(vehicle_model, "swift")
    plant_key = _PLANT_MAP.get(plant, "manesar")
    comp_key = _COMP_MAP.get(component, "engine_mount")
    folder = "drawings" if doc_type == "drawing" else "wis"

    # Primary filename: maruti_swift_manesar_engine_mount.json
    filepath = SEEDED_DATA_DIR / folder / f"maruti_{model_key}_{plant_key}_{comp_key}.json"

    if not filepath.exists():
        # Fallback: try any plant variant for this vehicle+component
        pattern = f"maruti_{model_key}_*_{comp_key}.json"
        matches = list((SEEDED_DATA_DIR / folder).glob(pattern))
        if matches:
            filepath = matches[0]

    if filepath.exists():
        with open(filepath) as f:
            data = json.load(f)
        # For WIS documents: strip drawing-only dimensional fields so they don't
        # create false deviations against the LLM-extracted drawing values.
        if doc_type == "wis":
            data = {k: v for k, v in data.items() if k not in _WIS_EXCLUDE_FIELDS}
        return data

    defaults = _generic_defaults(vehicle_model, component)
    if doc_type == "wis":
        defaults = {k: v for k, v in defaults.items() if k not in _WIS_EXCLUDE_FIELDS}
    return defaults


def _generic_defaults(vehicle_model: str, component: str) -> dict:
    return {
        "drawing_number": "MSIL-DE-SW-0001-C",
        "drawing_revision": "C",
        "part_number": "SW-GEN-001-C",
        "part_description": component,
        "vehicle_model": vehicle_model,
        "vehicle_variant": "VXI",
        "plant": "Manesar",
        "component": component,
        "bore_diameter_mm": 64.0,
        "bore_tolerance_mm": 0.02,
        "shaft_length_mm": 120.0,
        "gap_tolerance_mm": 0.15,
        "overall_width_mm": 180.0,
        "overall_height_mm": 95.0,
        "Ra_value": 1.6,
        "treatment": "Phosphate + Paint",
        "primary_torque_nm": 65.0,
        "fastener_grade": "8.8",
        "fastener_size": "M10",
        "torque_sequence": "cross-pattern",
        "re_torque_required": False,
        "material_grade": "IS 2062 Gr. B",
        "heat_treatment": "NIL",
        "thickness_mm": 6.0,
        "flatness_tolerance_mm": 0.10,
        "parallelism_tolerance_mm": 0.10,
        "true_position_dia_mm": 0.08,
        "datum_references": ["A", "B", "C"],
        "operation_step": 1,
        "sub_assembly_order": 1,
        "tooling_reference": "TL-001-NEW",
    }
